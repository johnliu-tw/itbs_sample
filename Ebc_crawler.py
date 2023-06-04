from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from bs4 import BeautifulSoup
import os
import traceback

import csv
from docx import Document

from datetime import date
today = date.today()
# 修改日期格式
search_date = today.strftime('%m/%d')
file_date = today.strftime('%Y-%m-%d')

# 設定無圖片與影片讀取模式
options = Options()
setting = {"profile.managed_default_content_settings.images": 2, "profile.managed_default_content_settings.video": 2}
options.add_experimental_option("prefs", setting)
# 開啟瀏覽器
service = Service(os.getcwd() + '/chromedriver')
driver = webdriver.Chrome(service=service, options=options)

try:
  # 開啟東森網站
  driver.get('https://news.ebc.net.tw/news/living')
  # 建立 Word 檔案
  document = Document()
  document.add_heading(file_date+' 東森資料', 0)
  # 建立 Csv 檔案
  csv_file = open(file_date+'_data.csv', 'a', newline='', encoding='utf_8_sig')
  writer = csv.writer(csv_file)
  for index in range(1, 4):
    sourceCode = BeautifulSoup(driver.page_source, "html.parser")

    # 爬取資料
    article_box = sourceCode.select('div.news-list-box')[0]
    articles = article_box.select('div.style1.white-box')
    for article in articles:
        date_string = article.select('span.small-gray-text')[0].text
        if search_date not in date_string:
          continue

        title = article.select('span.title')[0].text
        summary = article.select('span.summary')[0].text
        print(title)
        print(date_string)
        print(summary)

        # 寫入 Csv
        writer.writerow([title, date_string, summary])

        # 寫入 Word
        document.add_heading(title, 1)
        paragraph = document.add_paragraph('')
        paragraph.add_run(date_string).bold = True
        document.add_paragraph(summary)

    # 模擬動態換頁按鈕點擊
    element = driver.find_elements(By.CSS_SELECTOR, 'div.page-area a.white-btn')[-1]
    driver.execute_script("arguments[0].scrollIntoView();", element)
    element.click()

  # 關閉與儲存 csv + word 檔案
  csv_file.close()
  document.save(file_date+' 東森資料.docx')

except Exception as e:
  csv_file.close()
  traceback.print_exc()
  driver.close()